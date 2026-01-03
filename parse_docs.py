import os
from docx import Document
from docx.table import Table

def parse_docx(filepath):
    """Parse a DOCX file and extract all content including tables."""
    doc = Document(filepath)
    output = []
    output.append(f"\n{'='*80}")
    output.append(f"FILE: {os.path.basename(filepath)}")
    output.append(f"{'='*80}\n")
    
    for element in doc.element.body:
        if element.tag.endswith('p'):
            # Paragraph
            for para in doc.paragraphs:
                if para._element == element:
                    if para.text.strip():
                        style = para.style.name if para.style else "Normal"
                        output.append(f"[{style}] {para.text}")
                    break
        elif element.tag.endswith('tbl'):
            # Table
            for table in doc.tables:
                if table._tbl == element:
                    output.append("\n[TABLE START]")
                    for i, row in enumerate(table.rows):
                        cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
                        output.append(f"  Row {i}: {' | '.join(cells)}")
                    output.append("[TABLE END]\n")
                    break
    
    return '\n'.join(output)

# Parse all DOCX files
sample_dir = r"c:\Users\chara\Documents\QA Report App\sample"
docx_files = [
    "QA and Compliance Checklist 2026 BLANK.docx",
    "Level II Tier QA & Compliance Checklist_2026.docx",
    "Tyrone Fall_2025 with Summer Baseline _2025OverlayQA.docx"
]

with open("parsed_output.txt", "w", encoding="utf-8") as outfile:
    for docx_file in docx_files:
        filepath = os.path.join(sample_dir, docx_file)
        if os.path.exists(filepath):
            outfile.write(parse_docx(filepath))
        else:
            outfile.write(f"File not found: {docx_file}\n")

print("Parsing complete. Output saved to parsed_output.txt")
